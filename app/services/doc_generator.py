"""
DocumentGenerator — RF5.1, RF5.2, RF5.3, RF5.4

Genera CV su misura, cover letter e bozza email per ogni match.
L'enfasi è sempre sull'annuncio specifico: le stesse skill vengono
presentate diversamente a seconda di cosa cerca l'azienda.

Export in PDF (reportlab) e DOCX (python-docx).
"""
import json
import logging
import re
from pathlib import Path

from app.core.config import OUTPUT_DIR
from app.services.llm_gateway import LLMGateway

logger = logging.getLogger(__name__)


class DocumentGenerator:
    def __init__(self, llm: LLMGateway):
        self.llm = llm

    # -----------------------------------------------------------------
    async def generate(
        self,
        doc_type: str,
        user_profile: dict,
        listing: dict,
        match_info: dict | None = None,
    ) -> str:
        if doc_type == "cv":
            return await self._generate_cv(user_profile, listing, match_info)
        if doc_type == "cover_letter":
            return await self._generate_cover_letter(user_profile, listing, match_info)
        if doc_type == "email":
            return await self._generate_email(user_profile, listing, match_info)
        raise ValueError(f"doc_type non supportato: {doc_type}")

    # -----------------------------------------------------------------
    @staticmethod
    def _base_block(profile: dict, doc_type: str) -> str:
        """Il documento base dell'utente, se ne ha uno.

        Cambia la natura della richiesta all'LLM: senza, ogni CV viene
        riscritto da zero dai dati grezzi del profilo, e due CV per due
        aziende diverse possono raccontare la stessa persona in modi
        scollegati. Con il documento base, il compito diventa ADATTARE un
        testo già valido — più corto da chiedere e più coerente da
        leggere."""
        base = (profile.get("base_documents") or {}).get(doc_type)
        if not base:
            return ""
        return (
            f"\n=== {doc_type.upper()} BASE DEL CANDIDATO ===\n"
            "Questo è il documento generale già approvato dal candidato. "
            "ADATTALO all'annuncio qui sotto: riordina e riformula per "
            "mettere in evidenza ciò che l'annuncio chiede, ma non "
            "stravolgerlo e non aggiungere esperienze che non ci sono.\n\n"
            f"{base[:6000]}\n"
        )

    def _profile_block(self, p: dict) -> str:
        lines = [
            f"Nome: {p.get('name')}",
            f"Email: {p.get('email')}",
        ]
        if p.get("phone"):
            lines.append(f"Telefono: {p['phone']}")
        if p.get("location"):
            lines.append(f"Località: {p['location']}")
        if p.get("bio"):
            lines.append(f"Profilo: {p['bio']}")

        if p.get("hard_skills"):
            hs = ", ".join(
                f"{s['name']} ({s.get('level', 'n/d')}, {s.get('years_exp', 0)} anni)"
                for s in p["hard_skills"]
            )
            lines.append(f"Competenze tecniche: {hs}")

        if p.get("soft_skills"):
            lines.append(f"Soft skill: {', '.join(p['soft_skills'])}")

        if p.get("experiences"):
            lines.append("Esperienze lavorative:")
            for e in p["experiences"]:
                period = f"{e.get('start_date', '?')} – {e.get('end_date') or 'oggi'}"
                lines.append(f"  • {e.get('role')} presso {e.get('company')} ({period})")
                if e.get("description"):
                    lines.append(f"    {e['description']}")
                if e.get("skills_used"):
                    lines.append(f"    Tecnologie: {e['skills_used']}")

        if p.get("education"):
            lines.append("Formazione:")
            for ed in p["education"]:
                lines.append(
                    f"  • {ed.get('degree')} in {ed.get('field')} — "
                    f"{ed.get('institution')} ({ed.get('year')})"
                )
        return "\n".join(lines)

    def _listing_block(self, listing: dict) -> str:
        return (
            f"Posizione: {listing.get('title')}\n"
            f"Azienda: {listing.get('company_name')}\n"
            f"Sede: {listing.get('location') or 'non specificata'}\n"
            f"Requisiti:\n{(listing.get('requirements_raw') or '')[:2500]}"
        )

    @staticmethod
    def _match_hint(match_info: dict | None) -> str:
        if not match_info:
            return ""
        try:
            matched = json.loads(match_info.get("matched_skills") or "[]")
            missing = json.loads(match_info.get("missing_skills") or "[]")
        except (json.JSONDecodeError, TypeError):
            matched, missing = [], []

        parts = []
        if matched:
            parts.append(f"Punti di forza da enfatizzare: {', '.join(matched)}")
        if missing:
            parts.append(
                f"Competenze richieste non possedute (NON mentire, gestisci con onestà "
                f"mostrando capacità di apprendere): {', '.join(missing)}"
            )
        return "\n".join(parts)

    # -----------------------------------------------------------------
    async def _generate_cv(
        self, profile: dict, listing: dict, match_info: dict | None
    ) -> str:
        prompt = (
            "Scrivi un CV professionale in italiano, in formato Markdown, "
            "SU MISURA per l'annuncio indicato.\n\n"
            "REGOLE IMPORTANTI:\n"
            "- Usa SOLO informazioni presenti nel profilo. Non inventare NULLA.\n"
            "- Riordina esperienze e competenze mettendo per prime quelle rilevanti "
            "per questo annuncio specifico.\n"
            "- Riformula le descrizioni delle esperienze evidenziando gli aspetti "
            "che l'azienda cerca.\n"
            "- Struttura: intestazione con contatti, profilo breve, competenze, "
            "esperienze, formazione.\n"
            "- Massimo una pagina.\n\n"
            f"=== PROFILO CANDIDATO ===\n{self._profile_block(profile)}\n"
            f"{self._base_block(profile, 'cv')}\n"
            f"=== ANNUNCIO TARGET ===\n{self._listing_block(listing)}\n\n"
            f"{self._match_hint(match_info)}"
        )
        return await self.llm.complete(prompt)

    async def _generate_cover_letter(
        self, profile: dict, listing: dict, match_info: dict | None
    ) -> str:
        prompt = (
            "Scrivi una lettera di presentazione in italiano per questa candidatura.\n\n"
            "REGOLE:\n"
            "- Massimo 300 parole.\n"
            "- Tono professionale ma non ingessato.\n"
            "- Collega esplicitamente l'esperienza del candidato a ciò che l'annuncio chiede.\n"
            "- Non inventare esperienze o competenze non presenti nel profilo.\n"
            "- Niente frasi fatte tipo 'sono una persona dinamica e proattiva'.\n"
            "- Chiudi con una call to action concreta.\n\n"
            f"=== PROFILO CANDIDATO ===\n{self._profile_block(profile)}\n"
            f"{self._base_block(profile, 'cover_letter')}\n"
            f"=== ANNUNCIO TARGET ===\n{self._listing_block(listing)}\n\n"
            f"{self._match_hint(match_info)}"
        )
        return await self.llm.complete(prompt)

    async def _generate_email(
        self, profile: dict, listing: dict, match_info: dict | None
    ) -> str:
        prompt = (
            "Scrivi una email breve per inviare la candidatura a questa azienda.\n\n"
            "REGOLE:\n"
            "- Massimo 120 parole, deve essere leggibile in 20 secondi.\n"
            "- Include una riga oggetto efficace.\n"
            "- Formato: 'Oggetto: ...' seguito da riga vuota e poi il corpo.\n"
            "- Menziona che CV e lettera sono in allegato.\n"
            "- Non ripetere il contenuto della cover letter, sii complementare.\n\n"
            f"=== CANDIDATO ===\n{profile.get('name')} — {profile.get('email')}\n"
            f"{profile.get('bio') or ''}\n\n"
            f"=== ANNUNCIO ===\n{listing.get('title')} presso {listing.get('company_name')}"
        )
        return await self.llm.complete(prompt)

    # -----------------------------------------------------------------
    # RF5.4 — export
    # -----------------------------------------------------------------
    @staticmethod
    def output_stem(doc_type: str, company_name: str | None, title: str | None, doc_id: str) -> str:
        """Nome file deterministico condiviso da tutti i formati esportati
        di un documento (pdf/docx/md/txt). Ricalcolabile da chi deve
        ritrovare i file su disco (es. cleanup) senza salvare il path in DB.

        Gli ultimi 8 caratteri di `doc_id` (suffisso di unicità) non vengono
        MAI troncati: si accorcia prima la parte descrittiva, poi si
        appende sempre l'id per intero (71 + 1 + 8 = 80 caratteri max)."""
        id_suffix = re.sub(r"[^A-Za-z0-9_\-]", "_", doc_id[:8])
        descriptive = "_".join(p for p in (doc_type, company_name, title) if p)
        descriptive = re.sub(r"[^A-Za-z0-9_\-]", "_", descriptive)[:71]
        return f"{descriptive}_{id_suffix}" if descriptive else id_suffix

    @staticmethod
    def export(content: str, fmt: str, filename_stem: str) -> Path:
        """Esporta il documento nel formato richiesto e ritorna il path."""
        OUTPUT_DIR.mkdir(exist_ok=True, parents=True)

        if fmt == "md":
            path = OUTPUT_DIR / f"{filename_stem}.md"
            path.write_text(content, encoding="utf-8")
            return path

        if fmt == "txt":
            path = OUTPUT_DIR / f"{filename_stem}.txt"
            path.write_text(content, encoding="utf-8")
            return path

        if fmt == "pdf":
            return DocumentGenerator._export_pdf(content, filename_stem)

        if fmt == "docx":
            return DocumentGenerator._export_docx(content, filename_stem)

        raise ValueError(f"Formato non supportato: {fmt}")

    @staticmethod
    def _export_pdf(content: str, stem: str) -> Path:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        path = OUTPUT_DIR / f"{stem}.pdf"
        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )

        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_LEFT
        )
        h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, spaceAfter=6)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceAfter=4)

        story = []
        for raw in content.split("\n"):
            line = raw.rstrip()
            if not line.strip():
                story.append(Spacer(1, 4))
                continue
            if line.startswith("### "):
                story.append(Paragraph(line[4:], h2))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], h1))
            elif line.lstrip().startswith(("- ", "* ", "• ")):
                txt = line.lstrip()[2:]
                story.append(Paragraph(f"• {DocumentGenerator._md_inline(txt)}", body))
            else:
                story.append(Paragraph(DocumentGenerator._md_inline(line), body))

        doc.build(story)
        return path

    @staticmethod
    def _export_docx(content: str, stem: str) -> Path:
        from docx import Document
        from docx.shared import Pt

        path = OUTPUT_DIR / f"{stem}.docx"
        d = Document()
        style = d.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        for raw in content.split("\n"):
            line = raw.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                d.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                d.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                d.add_heading(line[2:], level=1)
            elif line.lstrip().startswith(("- ", "* ", "• ")):
                d.add_paragraph(line.lstrip()[2:].replace("**", ""), style="List Bullet")
            else:
                d.add_paragraph(line.replace("**", ""))

        d.save(str(path))
        return path

    @staticmethod
    def _md_inline(text: str) -> str:
        """Converte il grassetto markdown in tag reportlab."""
        out, bold = [], False
        i = 0
        while i < len(text):
            if text[i : i + 2] == "**":
                out.append("</b>" if bold else "<b>")
                bold = not bold
                i += 2
            else:
                ch = text[i]
                out.append({"&": "&amp;", "<": "&lt;", ">": "&gt;"}.get(ch, ch))
                i += 1
        if bold:
            out.append("</b>")
        return "".join(out)
